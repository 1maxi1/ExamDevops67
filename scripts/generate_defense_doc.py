# ruff: noqa: E501

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Защита_проекта_ExamDevops6_Билет22.docx"


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_subtitle(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(11)


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_code(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def build_document() -> Document:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_title(document, "Отчёт для защиты проекта ExamDevops6")
    add_subtitle(document, "Билет 22. Приложение мастер-классов гончарной мастерской")
    add_subtitle(document, "Дата подготовки: 13.04.2026")

    document.add_heading("1. Назначение проекта", level=1)
    document.add_paragraph(
        "Проект реализует веб-приложение по билету 22. Система предоставляет две "
        "основные функции: получение информации о наименованиях мастер-классов и "
        "получение информации о записи на мастер-класс по SMS-коду. Дополнительно "
        "вместе с мастер-классами возвращается их стоимость."
    )
    document.add_paragraph(
        "Архитектурно проект состоит из веб-сервера FastAPI, базы данных PostgreSQL "
        "для контейнерного запуска, альтернативной SQLite-конфигурации для тестов "
        "и одиночного deploy, графического интерфейса на HTML/CSS/JavaScript, "
        "Docker-обвязки и конфигураций для TeamCity."
    )

    document.add_heading("2. Структура проекта", level=1)
    for item in [
        "app/config.py — настройки приложения и чтение переменных окружения.",
        "app/database.py — создание engine, session factory и dependency для БД.",
        "app/models.py — SQLAlchemy-модели таблиц workshop_classes и workshop_registrations.",
        "app/schemas.py — Pydantic-схемы запросов и ответов API.",
        "app/seed.py — начальные данные, которыми заполняется база при старте.",
        "app/main.py — точка входа FastAPI, жизненный цикл и все HTTP-маршруты.",
        "app/static/index.html — графический интерфейс приложения.",
        "tests/conftest.py — тестовая конфигурация и клиент FastAPI.",
        "tests/test_api.py — автоматические тесты API и главной страницы.",
        "Dockerfile — сборка образа приложения.",
        "docker-compose.yml — запуск связки API + PostgreSQL.",
        ".teamcity/settings.kts — Kotlin DSL для TeamCity.",
    ]:
        add_bullet(document, item)

    document.add_heading("3. Подробный разбор Python-кода", level=1)

    document.add_heading("3.1 app/config.py", level=2)
    document.add_paragraph(
        "Файл содержит конфигурацию приложения. Класс Settings наследуется от "
        "BaseSettings из pydantic-settings, поэтому может читать значения как из "
        "переменных окружения, так и из файла .env."
    )
    add_bullet(document, "app_title — заголовок приложения.")
    add_bullet(document, "api_prefix — префикс REST API, сейчас это /api/v1.")
    add_bullet(
        document,
        "database_url — строка подключения к БД. По умолчанию указывает на PostgreSQL "
        "сервис db внутри docker-compose.",
    )
    add_bullet(document, "debug — флаг режима отладки FastAPI.")
    add_bullet(document, "seed_database — разрешение на автоматическое сидирование БД.")
    document.add_paragraph(
        "Функция get_settings() обёрнута в lru_cache, поэтому объект настроек "
        "создаётся один раз и затем переиспользуется."
    )

    document.add_heading("3.2 app/database.py", level=2)
    add_bullet(
        document,
        "create_engine_from_settings(settings) создаёт SQLAlchemy engine и "
        "добавляет check_same_thread=False для SQLite.",
    )
    add_bullet(
        document,
        "create_session_factory(engine) создаёт sessionmaker с управляемым поведением сессии.",
    )
    add_bullet(document, "init_db(engine) создаёт таблицы через Base.metadata.create_all.")
    add_bullet(
        document,
        "get_db_session(request) выдаёт сессию БД на время обработки HTTP-запроса.",
    )

    document.add_heading("3.3 app/models.py", level=2)
    add_bullet(document, "WorkshopClass — модель мастер-класса с полями id, name и price.")
    add_bullet(
        document,
        "WorkshopRegistration — модель записи на мастер-класс с полями id, sms_code, workshop_name и registration_time.",
    )

    document.add_heading("3.4 app/schemas.py", level=2)
    add_bullet(
        document,
        "WorkshopClassResponse описывает ответ для публичного списка мастер-классов.",
    )
    add_bullet(
        document,
        "RegistrationRequest описывает входной JSON с одним SMS-кодом.",
    )
    add_bullet(
        document,
        "RegistrationResponse описывает ответ по записи на мастер-класс.",
    )

    document.add_heading("3.5 app/seed.py", level=2)
    add_bullet(
        document,
        "DEFAULT_WORKSHOP_CLASSES хранит три записи мастер-классов из билета 22.",
    )
    add_bullet(
        document,
        "DEFAULT_WORKSHOP_REGISTRATIONS хранит три записи на мастер-классы из билета 22.",
    )
    add_bullet(
        document,
        "seed_initial_data(session) сначала проверяет пустоту таблиц, потом добавляет записи и делает commit.",
    )

    document.add_heading("3.6 app/main.py", level=2)
    add_bullet(document, "STATIC_DIR вычисляет путь к каталогу со статическим интерфейсом.")
    add_bullet(
        document,
        "create_app(settings=None) — фабрика приложения, позволяющая запускать его с разными настройками.",
    )
    add_bullet(
        document,
        "app_settings = settings or get_settings() — выбор переданных настроек или загрузка настроек по умолчанию.",
    )
    add_bullet(document, "engine и session_factory создаются один раз на приложение.")
    add_bullet(
        document,
        "lifespan(app) при старте создаёт таблицы и сидирует БД, а при завершении освобождает engine.",
    )
    add_bullet(document, "GET / отдаёт HTML-файл интерфейса.")
    add_bullet(document, "GET /health возвращает {\"status\": \"ok\"}.")
    add_bullet(
        document,
        "GET /api/v1/workshops возвращает список мастер-классов с названиями и стоимостью.",
    )
    add_bullet(
        document,
        "POST /api/v1/registrations/info принимает SMS-код, ищет запись в БД и возвращает информацию о записи или ошибку 401.",
    )
    add_bullet(
        document,
        "Последняя строка app = create_app() создаёт экспортируемый ASGI-объект для Uvicorn.",
    )

    document.add_heading("3.7 tests/conftest.py и tests/test_api.py", level=2)
    add_bullet(
        document,
        "Фикстура client создаёт временную SQLite-базу и TestClient.",
    )
    add_bullet(document, "test_healthcheck проверяет маршрут /health.")
    add_bullet(document, "test_index_page проверяет отдачу главной страницы.")
    add_bullet(document, "test_get_workshop_classes проверяет список мастер-классов.")
    add_bullet(document, "test_get_registration_success проверяет успешный поиск записи.")
    add_bullet(document, "test_get_registration_rejects_invalid_code проверяет отказ по неверному коду.")

    document.add_heading("4. Разбор интерфейса app/static/index.html", level=1)
    add_bullet(document, "Интерфейс встроен в один HTML-файл без отдельного фронтенд-сервера.")
    add_bullet(
        document,
        "Форма проверки записи отправляет POST-запрос на /api/v1/registrations/info.",
    )
    add_bullet(
        document,
        "Секция мастер-классов загружает публичный список через GET /api/v1/workshops.",
    )
    add_bullet(document, "loadHealth() показывает статус сервиса.")
    add_bullet(document, "renderWorkshops() отрисовывает карточки мастер-классов.")
    add_bullet(document, "renderSkeletons() показывает временные заглушки во время загрузки.")
    add_bullet(document, "setStatus() управляет сообщениями об успехе и ошибке.")

    document.add_heading("5. Dockerfile: подробный разбор", level=1)
    add_code(
        document,
        "FROM python:3.12-slim\n"
        "ENV PYTHONDONTWRITEBYTECODE=1\n"
        "ENV PYTHONUNBUFFERED=1\n"
        "WORKDIR /app\n"
        "COPY requirements.txt .\n"
        "RUN pip install --no-cache-dir -r requirements.txt\n"
        "COPY app ./app\n"
        "COPY app/static ./app/static\n"
        "EXPOSE 8000\n"
        "CMD [\"uvicorn\", \"app.main:app\", \"--host\", \"0.0.0.0\", \"--port\", \"8000\"]"
    )
    add_bullet(document, "FROM python:3.12-slim — базовый образ с Python 3.12.")
    add_bullet(
        document,
        "ENV PYTHONDONTWRITEBYTECODE=1 — Python не создаёт .pyc-файлы.",
    )
    add_bullet(
        document,
        "ENV PYTHONUNBUFFERED=1 — логи сразу выводятся в stdout/stderr.",
    )
    add_bullet(document, "WORKDIR /app — рабочая директория контейнера.")
    add_bullet(document, "COPY requirements.txt . — сначала копируется только файл зависимостей.")
    add_bullet(
        document,
        "RUN pip install --no-cache-dir -r requirements.txt — установка зависимостей приложения.",
    )
    add_bullet(document, "COPY app ./app — копирование исходного кода приложения.")
    add_bullet(
        document,
        "COPY app/static ./app/static — дублирующая инструкция, потому что папка static уже входит в COPY app ./app.",
    )
    add_bullet(document, "EXPOSE 8000 — документирует рабочий порт приложения.")
    add_bullet(document, "CMD [...] — запускает Uvicorn и публикует сервис на 0.0.0.0:8000.")
    document.add_paragraph(
        "Dockerfile написан в таком порядке ради кэширования. Сначала копируется "
        "requirements.txt и устанавливаются зависимости, а код приложения копируется "
        "позже. Если меняются только исходники, Docker может переиспользовать "
        "кэшированный слой с установленными пакетами."
    )
    document.add_paragraph(
        "Если говорить о стадиях сборки этого Dockerfile, практически значимыми шагами "
        "являются FROM, WORKDIR, COPY requirements.txt, RUN pip install, COPY app и "
        "COPY app/static. Базовый образ python:3.12-slim уже содержит собственные "
        "слои, поверх которых добавляются слои проекта."
    )

    document.add_heading("6. docker-compose.yml и volumes", level=1)
    add_bullet(
        document,
        "Сервис db использует образ postgres:16-alpine и хранит БД pottery_workshop.",
    )
    add_bullet(document, "Сервис api собирается из Dockerfile текущего проекта.")
    add_bullet(
        document,
        "api зависит от db через depends_on с condition: service_healthy.",
    )
    add_bullet(
        document,
        "DATABASE_URL использует хост db — это имя сервиса внутри сети Docker Compose.",
    )
    add_bullet(document, "Порт 5432 пробрасывается для БД, а 8000 — для API.")
    add_bullet(
        document,
        "Volume postgres_data монтируется в /var/lib/postgresql/data и сохраняет данные базы между перезапусками.",
    )

    document.add_heading("7. TeamCity и CI/CD", level=1)
    add_bullet(
        document,
        "Для main-ветки собирается Docker-образ, отправляется в DockerHub и запускается на prod-стенде.",
    )
    add_bullet(
        document,
        "Для feature/fix-веток собирается Docker-образ без deploy.",
    )
    add_bullet(document, "Ruff используется как линтер.")
    add_bullet(document, "Bandit используется как SAST-анализатор.")
    add_bullet(document, "TruffleHog используется для поиска секретов.")

    document.add_heading("8. Итог", level=1)
    document.add_paragraph(
        "Проект полностью переделан под билет 22 и покрывает серверную часть, базу данных, "
        "графический интерфейс, контейнеризацию и CI/CD-сценарии. Архитектура проекта "
        "осталась достаточно простой, чтобы её можно было быстро и уверенно объяснить на защите."
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("9. Краткая шпаргалка для защиты", level=1)
    add_bullet(document, "Приложение состоит из FastAPI-сервера и PostgreSQL-базы данных.")
    add_bullet(document, "Публичная функция — список мастер-классов с ценами.")
    add_bullet(document, "Защищённая функция — запись на мастер-класс по SMS-коду.")
    add_bullet(
        document,
        "В Dockerfile сначала устанавливаются зависимости, потом копируется код — это сделано ради кэширования.",
    )
    add_bullet(
        document,
        "docker-compose поднимает два сервиса: api и db, а volume postgres_data сохраняет БД.",
    )
    add_bullet(
        document,
        "main-ветка в TeamCity выполняет build + push + deploy, feature/fix-ветки — build + проверки.",
    )

    return document


if __name__ == "__main__":
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)
